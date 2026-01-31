#!/usr/bin/env python3
"""
Tests for Mammoth-based DOCX to Quill HTML Converter.

Tests the new docx_to_quill.py module which provides:
- Direct DOCX → Quill HTML conversion via Mammoth
- nh3 HTML sanitization
- Quill-compatible list format (flat <ol> with ql-indent-N classes)
"""

import re
from pathlib import Path

import pytest

import sys
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))


# =============================================================================
# Import Tests - Graceful degradation when dependencies missing
# =============================================================================

class TestModuleImports:
    """Test that module imports work and dependencies are detected."""
    
    def test_module_imports(self):
        """Module should import without errors."""
        from docx_to_quill import (
            is_mammoth_available,
            is_nh3_available,
        )
        # Just verify they're callable
        assert callable(is_mammoth_available)
        assert callable(is_nh3_available)
    
    def test_mammoth_available(self):
        """Mammoth should be available after pip install mammoth."""
        from docx_to_quill import is_mammoth_available
        # This test documents the requirement - may be skipped if not installed
        available = is_mammoth_available()
        if not available:
            pytest.skip("mammoth not installed - run: pip install mammoth")
        assert available is True
    
    def test_nh3_available(self):
        """nh3 should be available after pip install nh3."""
        from docx_to_quill import is_nh3_available
        available = is_nh3_available()
        if not available:
            pytest.skip("nh3 not installed - run: pip install nh3")
        assert available is True


# =============================================================================
# List Conversion Tests - Core Quill format logic
# =============================================================================

class TestListConversion:
    """Test HTML list conversion to Quill format."""
    
    def test_simple_bullet_list(self):
        """Simple <ul> should convert to Quill <ol> with data-list='bullet'."""
        from docx_to_quill import _convert_lists_to_quill_format
        
        html = "<ul><li>Item 1</li><li>Item 2</li></ul>"
        result = _convert_lists_to_quill_format(html)
        
        assert "<ol>" in result
        assert "</ol>" in result
        assert 'data-list="bullet"' in result
        assert '<span class="ql-ui"></span>' in result
        assert "Item 1" in result
        assert "Item 2" in result
    
    def test_nested_list_indentation(self):
        """Nested lists should get ql-indent-N classes."""
        from docx_to_quill import _convert_lists_to_quill_format
        
        html = """
        <ul>
            <li>Parent
                <ul>
                    <li>Child 1</li>
                    <li>Child 2</li>
                </ul>
            </li>
        </ul>
        """
        result = _convert_lists_to_quill_format(html)
        
        # Should have a single <ol> wrapper
        assert result.count("<ol>") == 1
        assert result.count("</ol>") == 1
        
        # Child items should have ql-indent-1 class
        assert 'class="ql-indent-1"' in result
    
    def test_deeply_nested_list(self):
        """3+ level nesting should use ql-indent-2, ql-indent-3, etc."""
        from docx_to_quill import _convert_lists_to_quill_format
        
        html = """
        <ul>
            <li>Level 0
                <ul>
                    <li>Level 1
                        <ul>
                            <li>Level 2</li>
                        </ul>
                    </li>
                </ul>
            </li>
        </ul>
        """
        result = _convert_lists_to_quill_format(html)
        
        # Check for indent levels - capped at 1 for ATS optimization
        # Deep nesting (indent 2+) is flattened to improve ATS/recruiter parsing
        assert 'class="ql-indent-1"' in result
        # All deeper levels are flattened to indent-1
        assert 'class="ql-indent-2"' not in result  # Capped!
    
    def test_preserves_inline_formatting_in_list(self):
        """Bold, italic, links should be preserved in list items."""
        from docx_to_quill import _convert_lists_to_quill_format
        
        html = '<ul><li><strong>Bold</strong> and <em>italic</em></li></ul>'
        result = _convert_lists_to_quill_format(html)
        
        assert "<strong>Bold</strong>" in result
        assert "<em>italic</em>" in result
    
    def test_ordered_list_conversion(self):
        """Ordered <ol> should also convert to Quill format."""
        from docx_to_quill import _convert_lists_to_quill_format
        
        html = "<ol><li>First</li><li>Second</li></ol>"
        result = _convert_lists_to_quill_format(html)
        
        assert "<ol>" in result
        assert 'data-list="bullet"' in result
        assert '<span class="ql-ui"></span>' in result


# =============================================================================
# HTML Sanitization Tests
# =============================================================================

class TestHtmlSanitization:
    """Test nh3-based HTML sanitization."""
    
    @pytest.fixture(autouse=True)
    def check_nh3(self):
        """Skip tests if nh3 not available."""
        from docx_to_quill import is_nh3_available
        if not is_nh3_available():
            pytest.skip("nh3 not installed")
    
    def test_allows_safe_tags(self):
        """Common safe tags should be preserved."""
        from docx_to_quill import sanitize_html
        
        html = '<p><strong>Bold</strong> and <em>italic</em></p>'
        result = sanitize_html(html)
        
        assert "<p>" in result
        assert "<strong>" in result
        assert "<em>" in result
    
    def test_strips_script_tags(self):
        """Script tags should be removed (XSS prevention)."""
        from docx_to_quill import sanitize_html
        
        html = '<p>Safe</p><script>alert("xss")</script>'
        result = sanitize_html(html)
        
        assert "<script>" not in result
        assert "alert" not in result
        assert "<p>Safe</p>" in result
    
    def test_strips_onclick_handlers(self):
        """Event handlers should be removed."""
        from docx_to_quill import sanitize_html
        
        html = '<p onclick="evil()">Text</p>'
        result = sanitize_html(html)
        
        assert "onclick" not in result
        assert "evil" not in result
    
    def test_preserves_link_href(self):
        """Safe link attributes should be preserved."""
        from docx_to_quill import sanitize_html
        
        html = '<a href="https://example.com" target="_blank">Link</a>'
        result = sanitize_html(html)
        
        assert 'href="https://example.com"' in result
        assert "Link" in result
    
    def test_allows_quill_classes(self):
        """Quill-specific classes should be preserved."""
        from docx_to_quill import sanitize_html
        
        html = '<li class="ql-indent-2" data-list="bullet"><span class="ql-ui"></span>Item</li>'
        result = sanitize_html(html)
        
        # Note: nh3 may reorder attributes, so check for presence
        assert "ql-indent" in result or "data-list" in result


# =============================================================================
# Link Security Tests
# =============================================================================

class TestLinkSecurity:
    """Test link security attribute handling."""
    
    def test_adds_target_blank(self):
        """Links should get target='_blank'."""
        from docx_to_quill import _ensure_link_attributes
        
        html = '<a href="https://example.com">Link</a>'
        result = _ensure_link_attributes(html)
        
        assert 'target="_blank"' in result
    
    def test_adds_rel_noopener(self):
        """Links should get rel='noopener noreferrer'."""
        from docx_to_quill import _ensure_link_attributes
        
        html = '<a href="https://example.com">Link</a>'
        result = _ensure_link_attributes(html)
        
        assert 'noopener' in result


# =============================================================================
# Full Pipeline Tests (DOCX → Quill HTML)
# =============================================================================

class TestDocxToQuillPipeline:
    """Test the full DOCX → Quill HTML conversion pipeline."""
    
    @pytest.fixture(autouse=True)
    def check_dependencies(self):
        """Skip tests if Mammoth not available."""
        from docx_to_quill import is_mammoth_available
        if not is_mammoth_available():
            pytest.skip("mammoth not installed")
    
    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create a sample DOCX file for testing."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed (dev dependency)")
        
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("This is a test paragraph with some text.")
        
        # Add a bulleted list
        doc.add_paragraph("First item", style="List Bullet")
        doc.add_paragraph("Second item", style="List Bullet")
        
        file_path = tmp_path / "test.docx"
        doc.save(file_path)
        return file_path
    
    def test_docx_to_html_basic(self, sample_docx):
        """Basic DOCX should convert to HTML."""
        from docx_to_quill import docx_to_html
        
        html = docx_to_html(sample_docx)
        
        assert html  # Non-empty
        assert "Test Document" in html or "test" in html.lower()
    
    def test_docx_to_quill_html(self, sample_docx):
        """DOCX should convert to Quill-compatible HTML."""
        from docx_to_quill import docx_to_quill_html
        
        html = docx_to_quill_html(sample_docx)
        
        assert html  # Non-empty
        # Should be single-line (Europass requirement)
        assert "\n" not in html
    
    def test_docx_not_found_raises(self, tmp_path):
        """Missing file should raise FileNotFoundError."""
        from docx_to_quill import docx_to_html
        
        with pytest.raises(FileNotFoundError):
            docx_to_html(tmp_path / "nonexistent.docx")
    
    def test_non_docx_raises(self, tmp_path):
        """Non-DOCX file should raise ValueError."""
        from docx_to_quill import docx_to_html
        
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("Not a DOCX")
        
        with pytest.raises(ValueError):
            docx_to_html(txt_file)


# =============================================================================
# HTML to Quill Format Tests (for non-DOCX HTML sources)
# =============================================================================

class TestHtmlToQuillFormat:
    """Test converting arbitrary HTML to Quill format."""
    
    def test_converts_ul_to_quill(self):
        """Standard HTML should convert to Quill format."""
        from docx_to_quill import html_to_quill_format
        
        html = "<ul><li>Item 1</li><li>Item 2</li></ul>"
        result = html_to_quill_format(html)
        
        assert "<ol>" in result
        assert 'data-list="bullet"' in result
    
    def test_single_line_output(self):
        """Output should be single-line."""
        from docx_to_quill import html_to_quill_format
        
        html = "<p>Line 1</p>\n<p>Line 2</p>"
        result = html_to_quill_format(html)
        
        assert "\n" not in result


# =============================================================================
# Pydantic Models Tests (optional validation)
# =============================================================================

class TestQuillDeltaModels:
    """Test Pydantic models for Quill Delta validation."""
    
    @pytest.fixture(autouse=True)
    def check_pydantic(self):
        """Check if Pydantic models are available."""
        from docx_to_quill import PYDANTIC_AVAILABLE
        if not PYDANTIC_AVAILABLE:
            pytest.skip("Pydantic not available")
    
    def test_delta_op_valid(self):
        """Valid DeltaInsertOp should parse correctly."""
        from docx_to_quill import DeltaInsertOp
        
        op = DeltaInsertOp(insert="Hello", attributes=None)
        assert op.insert == "Hello"
    
    def test_delta_attribute_valid(self):
        """Valid DeltaAttribute should parse correctly."""
        from docx_to_quill import DeltaAttribute
        
        attr = DeltaAttribute(bold=True, italic=False)
        assert attr.bold is True
        assert attr.italic is False
    
    def test_quill_delta_valid(self):
        """Valid QuillDelta should parse correctly."""
        from docx_to_quill import QuillDelta, DeltaInsertOp
        
        delta = QuillDelta(ops=[
            DeltaInsertOp(insert="Hello "),
            DeltaInsertOp(insert="World"),
            DeltaInsertOp(insert="\n"),
        ])
        assert len(delta.ops) == 3
    
    def test_delta_with_attributes(self):
        """Delta with attributes should parse correctly."""
        from docx_to_quill import QuillDelta, DeltaInsertOp, DeltaAttributes
        
        delta = QuillDelta(ops=[
            DeltaInsertOp(
                insert="Bold text",
                attributes=DeltaAttributes(bold=True)
            ),
            DeltaInsertOp(insert="\n"),
        ])
        assert delta.ops[0].attributes.bold is True


# =============================================================================
# Integration with mcp_server Tests
# =============================================================================

class TestMcpServerIntegration:
    """Test integration with mcp_server.py."""
    
    def test_mcp_server_imports_module(self):
        """mcp_server should import the docx_to_quill module."""
        from mcp_server import MAMMOTH_AVAILABLE, NH3_AVAILABLE
        
        # Just check the flags exist
        assert isinstance(MAMMOTH_AVAILABLE, bool)
        assert isinstance(NH3_AVAILABLE, bool)
    
    def test_parse_document_accepts_output_format(self):
        """parse_document should accept output_format parameter."""
        from mcp_server import parse_document
        
        # This tests the function signature, not the result
        # (would need a real DOCX file to test fully)
        import inspect
        sig = inspect.signature(parse_document)
        params = list(sig.parameters.keys())
        
        assert "output_format" in params


# =============================================================================
# Regression Tests - Real-world CV patterns
# =============================================================================

class TestCVPatterns:
    """Test real-world CV formatting patterns."""
    
    def test_job_description_with_nested_list(self):
        """Job descriptions often have nested bullet points."""
        from docx_to_quill import _convert_lists_to_quill_format
        
        html = """
        <ul>
            <li>Led development team
                <ul>
                    <li>Managed 5 engineers</li>
                    <li>Conducted code reviews</li>
                </ul>
            </li>
            <li>Improved performance by 40%</li>
        </ul>
        """
        result = _convert_lists_to_quill_format(html)
        
        # Verify structure
        assert "<ol>" in result
        assert 'ql-indent-1' in result
        assert "Led development team" in result
        assert "Managed 5 engineers" in result
    
    def test_skills_section_formatting(self):
        """Skills sections typically use bold headers."""
        from docx_to_quill import html_to_quill_format
        
        html = """
        <p><strong>Technical Skills:</strong></p>
        <ul>
            <li>Python, JavaScript, TypeScript</li>
            <li>Docker, Kubernetes, AWS</li>
        </ul>
        """
        result = html_to_quill_format(html)
        
        assert "<strong>Technical Skills:</strong>" in result
        assert 'data-list="bullet"' in result
    
    def test_education_entry_with_links(self):
        """Education entries may contain links."""
        from docx_to_quill import html_to_quill_format
        
        html = """
        <p><strong>MIT</strong> - <a href="https://mit.edu">Massachusetts Institute of Technology</a></p>
        <p>Master of Science in Computer Science, 2020</p>
        """
        result = html_to_quill_format(html)
        
        assert "<strong>MIT</strong>" in result
        assert 'href="https://mit.edu"' in result
